"""Compliance calendar routes (M4). No date math here — evaluator only (C12)."""
from __future__ import annotations

from typing import Any

import io
import uuid

from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import Response
from openpyxl import Workbook
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app import audit
from app.db import get_session
from app.models import (
    RuleExtension,
    Company,
    CalendarRow,
    CompanyFyAttributes,
    Role,
    RowStatus,
    User,
)
from app.repositories import calendar as cal_repo
from app.repositories import companies as companies_repo
from app.schemas.calendar import CalendarRowOut, FyAttributesIn, RowPatch
from app.security.auth import require_role
from app.services import calendar as svc

router = APIRouter(prefix="/api/v1", tags=["calendar"])

XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


async def _owned_company(session: AsyncSession, user: User, company_id: uuid.UUID) -> Company:
    company = await companies_repo.get_company(session, user.firm_id, company_id)
    if company is None:
        raise HTTPException(status_code=404, detail="company not found")
    return company


def _row_out(
    row: CalendarRow,
    code: str,
    version_no: int,
    payload: dict[str, Any],
    ext: RuleExtension | None,
) -> CalendarRowOut:
    extension_date = ext.extended_due_date if ext else None
    return CalendarRowOut(
        id=row.id,
        fy=row.fy,
        category=payload["category"],
        obligation_name=payload["obligation_name"],
        form_number=payload.get("form_number"),
        rule_code=code,
        rule_version=version_no,
        citation=payload["source_citation"],
        occurrence_label=row.occurrence_label,
        subject_type=row.subject_type.value,
        subject_id=row.subject_id,
        computed_due_date=row.computed_due_date,
        override_date=row.override_date,
        override_reason=row.override_reason,
        extension_date=extension_date,
        extension_ref=ext.circular_ref if ext else None,
        effective_due_date=row.override_date or extension_date or row.computed_due_date,
        status=row.status,
        srn=row.srn,
        filed_offline_ack=row.filed_offline_ack,
        assignee_user_id=row.assignee_user_id,
        remarks=row.remarks,
        needs_review=row.needs_review,
        needs_review_reason=row.needs_review_reason.value if row.needs_review_reason else None,
    )


@router.post("/companies/{company_id}/calendar/generate")
async def generate_calendar(
    company_id: uuid.UUID,
    fy: int,
    request: Request,
    user: User = Depends(require_role(Role.executive)),
    session: AsyncSession = Depends(get_session),
) -> dict[str, int]:
    company = await _owned_company(session, user, company_id)
    return await svc.generate(session, user.firm_id, company, fy, user.id, request)


@router.get("/companies/{company_id}/calendar", response_model=list[CalendarRowOut])
async def list_calendar(
    company_id: uuid.UUID,
    fy: int,
    needs_review: bool | None = None,
    user: User = Depends(require_role(Role.viewer)),
    session: AsyncSession = Depends(get_session),
) -> list[CalendarRowOut]:
    await _owned_company(session, user, company_id)
    rows = await cal_repo.rows_with_trace(
        session, user.firm_id, company_id, fy, needs_review=needs_review
    )
    return [_row_out(*r) for r in rows]


@router.patch("/calendar-rows/{row_id}")
async def patch_row(
    row_id: uuid.UUID,
    body: RowPatch,
    request: Request,
    user: User = Depends(require_role(Role.executive)),
    session: AsyncSession = Depends(get_session),
) -> dict[str, str]:
    row = await cal_repo.get_row(session, user.firm_id, row_id)
    if row is None:
        raise HTTPException(status_code=404, detail="calendar row not found")

    # PRD §9: Executive may edit assigned rows only
    if user.role == Role.executive and row.assignee_user_id != user.id:
        raise HTTPException(status_code=403, detail="executives may edit assigned rows only")

    fields = body.model_dump(exclude_unset=True)

    # PRD §9: override/extend is Manager+
    if ("override_date" in fields or "override_reason" in fields) and user.role == Role.executive:
        raise HTTPException(status_code=403, detail="overrides require Manager or Partner role")
    if "override_date" in fields and fields["override_date"] is not None:
        if not fields.get("override_reason") and not row.override_reason:
            raise HTTPException(status_code=422, detail="override requires a reason")

    # PRD §4.5: "Filed" requires an SRN or an explicit filed-offline ack
    new_status = fields.get("status", row.status)
    if new_status == RowStatus.filed:
        srn = fields.get("srn", row.srn)
        ack = fields.get("filed_offline_ack", row.filed_offline_ack)
        if not srn and not ack:
            raise HTTPException(
                status_code=422,
                detail="marking filed requires an SRN or the filed-offline acknowledgment",
            )

    before, after = {}, {}
    acknowledge = fields.pop("acknowledge_review", None)
    for key, value in fields.items():
        current = getattr(row, key)
        if current != value:
            before[key], after[key] = str(current), str(value)
            setattr(row, key, value)
    if acknowledge:
        before["needs_review"], after["needs_review"] = str(row.needs_review), "False"
        row.needs_review = False
        row.needs_review_reason = None

    if after:
        await audit.record(
            session, firm_id=user.firm_id, actor_user_id=user.id, entity_type="calendar_row",
            entity_id=row.id, action="update", before=before, after=after, request=request,
        )
        await session.commit()
    return {"status": row.status.value}


@router.get("/companies/{company_id}/calendar/export")
async def export_calendar(
    company_id: uuid.UUID,
    fy: int,
    user: User = Depends(require_role(Role.viewer)),
    session: AsyncSession = Depends(get_session),
) -> Response:
    await _owned_company(session, user, company_id)
    rows = await cal_repo.rows_with_trace(session, user.firm_id, company_id, fy)
    wb = Workbook()
    ws = wb.active
    ws.title = f"FY {fy - 1}-{str(fy)[2:]}"
    headers = [
        "category", "obligation", "form", "rule", "version", "citation", "occurrence",
        "computed_due_date", "extension_date", "override_date", "effective_due_date",
        "status", "srn", "needs_review", "remarks",
    ]
    ws.append(headers)
    for r in rows:
        out = _row_out(*r)
        ws.append([
            out.category, out.obligation_name, out.form_number, out.rule_code,
            out.rule_version, out.citation, out.occurrence_label,
            out.computed_due_date, out.extension_date, out.override_date,
            out.effective_due_date, out.status.value, out.srn, out.needs_review,
            out.remarks,
        ])
    buf = io.BytesIO()
    wb.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type=XLSX,
        headers={"Content-Disposition": f'attachment; filename="calendar_fy{fy}.xlsx"'},
    )


@router.get("/companies/{company_id}/calendar/export-word")
async def export_calendar_word(
    company_id: uuid.UUID,
    fy: int,
    user: User = Depends(require_role(Role.viewer)),
    session: AsyncSession = Depends(get_session),
) -> Response:
    """Word export of the calendar per entity (PRD §4.5). Stamped with
    generation time and rule versions — export fidelity per PRD §10."""
    company = await _owned_company(session, user, company_id)
    rows = await cal_repo.rows_with_trace(session, user.firm_id, company_id, fy)

    from datetime import datetime, timezone

    from docx import Document

    doc = Document()
    doc.add_heading(f"Compliance Calendar — {company.name}", level=1)
    doc.add_paragraph(f"CIN: {company.cin} · FY {fy - 1}-{str(fy)[2:]}")
    table = doc.add_table(rows=1, cols=7)
    header = table.rows[0].cells
    for i, title in enumerate(
        ["Category", "Obligation", "Form", "Due date", "Status", "Rule (version)", "Citation"]
    ):
        header[i].text = title
    for r in rows:
        out = _row_out(*r)
        cells = table.add_row().cells
        cells[0].text = out.category
        cells[1].text = out.obligation_name + (
            f" [{out.occurrence_label}]" if out.occurrence_label else ""
        )
        cells[2].text = out.form_number or "—"
        cells[3].text = str(out.effective_due_date or "needs review")
        cells[4].text = out.status.value + (" ⚑" if out.needs_review else "")
        cells[5].text = f"{out.rule_code} (v{out.rule_version})"
        cells[6].text = out.citation
    doc.add_paragraph(
        f"Generated {datetime.now(timezone.utc).isoformat(timespec='seconds')} — every "
        "date traces to the versioned rules dataset shown per row."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="calendar_fy{fy}.docx"'},
    )


@router.get("/companies/{company_id}/fy-attributes/{fy}")
async def get_fy_attributes(
    company_id: uuid.UUID,
    fy: int,
    user: User = Depends(require_role(Role.viewer)),
    session: AsyncSession = Depends(get_session),
) -> dict[str, Any]:
    await _owned_company(session, user, company_id)
    row = (
        await session.execute(
            select(CompanyFyAttributes).where(
                CompanyFyAttributes.firm_id == user.firm_id,
                CompanyFyAttributes.company_id == company_id,
                CompanyFyAttributes.fy == fy,
            )
        )
    ).scalar_one_or_none()
    fields = FyAttributesIn.model_fields.keys()
    if row is None:
        return {k: None for k in fields}
    return {k: (float(v) if k in ("turnover", "net_worth", "net_profit") and v is not None
                else v) for k, v in ((k, getattr(row, k)) for k in fields)}


@router.put("/companies/{company_id}/fy-attributes/{fy}")
async def put_fy_attributes(
    company_id: uuid.UUID,
    fy: int,
    body: FyAttributesIn,
    request: Request,
    user: User = Depends(require_role(Role.manager)),  # Manager+ (Amendment A1)
    session: AsyncSession = Depends(get_session),
) -> dict[str, Any]:
    await _owned_company(session, user, company_id)
    row = (
        await session.execute(
            select(CompanyFyAttributes).where(
                CompanyFyAttributes.firm_id == user.firm_id,
                CompanyFyAttributes.company_id == company_id,
                CompanyFyAttributes.fy == fy,
            )
        )
    ).scalar_one_or_none()
    data = body.model_dump(exclude_unset=True)
    if row is None:
        row = CompanyFyAttributes(firm_id=user.firm_id, company_id=company_id, fy=fy, **data)
        session.add(row)
        action, before = "create", None
    else:
        before = {k: str(getattr(row, k)) for k in data}
        for k, v in data.items():
            setattr(row, k, v)
        action = "update"
    await session.flush()
    await audit.record(
        session, firm_id=user.firm_id, actor_user_id=user.id, entity_type="company_fy_attributes",
        entity_id=row.id, action=action, before=before,
        after={k: str(v) for k, v in data.items()}, request=request,
    )
    await session.commit()
    return {"saved": True, "fy": fy}
