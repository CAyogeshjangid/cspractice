"""One-off: build the Phase 1 .docx template skeletons (charter C6: lives in
/scripts, never imported by app code).

Content is drafted in-house from the bare statutory skeleton (Companies Act
2013 / Secretarial Standards headings) — no competitor text, layouts, or
branding (charter 10.7). Templates ship UNSTAMPED: manifest.json carries
validated_* as null and generation refuses them until a professional stamps
each via the API (PRD §4.7).

Run from repo root:  python scripts/build_templates.py
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document

OUT = Path(__file__).resolve().parents[1] / "templates" / "docx"

LETTERHEAD_BLOCK = [
    "{% if letterhead_name %}{{ letterhead_name }}{% endif %}",
    "{% if letterhead_address %}{{ letterhead_address }}{% endif %}",
]

SIGNING_BLOCK = [
    "",
    "By order of the Board",
    "For {{ company_name }}",
    "",
    "_______________________",
    "{{ signatory_name }}",
    "{{ signatory_designation }}",
    "Date: {{ document_date }}",
    "Place: {{ place }}",
]


def build(filename: str, title: str, body: list[str]) -> None:
    doc = Document()
    for line in LETTERHEAD_BLOCK:
        doc.add_paragraph(line)
    heading = doc.add_heading(title, level=1)
    heading.alignment = 1  # centered
    for line in body:
        doc.add_paragraph(line)
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUT / filename)
    print(f"built {filename}")


build(
    "agm_notice.docx",
    "NOTICE OF ANNUAL GENERAL MEETING",
    [
        "NOTICE is hereby given that the Annual General Meeting of "
        "{{ company_name }} (CIN: {{ cin }}) will be held on {{ agm_date }} "
        "at {{ meeting_time }} at {{ venue }} to transact the following business:",
        "",
        "ORDINARY BUSINESS:",
        "{% for item in ordinary_business %}",
        "{{ loop.index }}. {{ item }}",
        "{% endfor %}",
        "{% if special_business %}",
        "SPECIAL BUSINESS:",
        "{% for item in special_business %}",
        "{{ loop.index }}. {{ item }}",
        "{% endfor %}",
        "{% endif %}",
        *SIGNING_BLOCK,
        "",
        "Notes: A member entitled to attend and vote is entitled to appoint a "
        "proxy to attend and vote instead of the member.",
    ],
)

build(
    "directors_report.docx",
    "BOARD'S REPORT",
    [
        "To the Members of {{ company_name }},",
        "Your Directors present their report together with the audited financial "
        "statements for the financial year ended {{ fy_end_date }}.",
        "",
        "1. FINANCIAL RESULTS (₹)",
        "Revenue from operations: {{ revenue }}",
        "Profit/(Loss) before tax: {{ profit_before_tax }}",
        "Profit/(Loss) after tax: {{ profit_after_tax }}",
        "",
        "2. STATE OF THE COMPANY'S AFFAIRS",
        "{{ state_of_affairs }}",
        "",
        "3. DIVIDEND",
        "{{ dividend }}",
        "",
        "4. TRANSFER TO RESERVES",
        "{{ transfer_to_reserves }}",
        "",
        "5. DIRECTORS AND KEY MANAGERIAL PERSONNEL",
        "{% for d in directors %}",
        "{{ d.name }} ({{ d.din }}) — {{ d.designation }}",
        "{% endfor %}",
        "",
        "6. NUMBER OF BOARD MEETINGS",
        "{{ board_meetings_held }}",
        "",
        "7. DIRECTORS' RESPONSIBILITY STATEMENT",
        "Pursuant to Section 134(5) of the Companies Act, 2013, the Board of "
        "Directors, to the best of their knowledge and ability, confirm the "
        "matters stated therein.",
        "",
        "8. DECLARATIONS AND DISCLOSURES",
        "{{ other_disclosures }}",
        *SIGNING_BLOCK,
    ],
)

build(
    "meeting_minutes.docx",
    "MINUTES OF THE MEETING",
    [
        "Minutes of the {{ meeting_type }} of {{ company_name }} "
        "(CIN: {{ cin }}) held on {{ meeting_date }} at {{ meeting_time }} "
        "at {{ venue }}.",
        "",
        "PRESENT:",
        "{% for d in directors %}",
        "{{ d.name }} — {{ d.designation }}",
        "{% endfor %}",
        "",
        "CHAIRPERSON: {{ chairperson }}",
        "",
        "PROCEEDINGS:",
        "{% for item in agenda_items %}",
        "{{ loop.index }}. {{ item }}",
        "{% endfor %}",
        "",
        "There being no other business, the meeting concluded at "
        "{{ conclusion_time }} with a vote of thanks to the Chair.",
        *SIGNING_BLOCK,
    ],
)

build(
    "attendance_sheet.docx",
    "ATTENDANCE SHEET",
    [
        "{{ meeting_type }} of {{ company_name }} (CIN: {{ cin }}) held on "
        "{{ meeting_date }} at {{ venue }}.",
        "",
        "ATTENDEES:",
        "{% for d in directors %}",
        "{{ loop.index }}. {{ d.name }} — {{ d.designation }} — DIN: {{ d.din }} — Signature: ______________",
        "{% endfor %}",
        *SIGNING_BLOCK,
    ],
)

build(
    "list_of_directors.docx",
    "LIST OF DIRECTORS",
    [
        "{{ company_name }} (CIN: {{ cin }}) — as on {{ document_date }}",
        "",
        "{% for d in directors %}",
        "{{ loop.index }}. {{ d.name }} — DIN: {{ d.din }} — {{ d.designation }} — "
        "Appointed: {{ d.appointment_date }}",
        "{% endfor %}",
        *SIGNING_BLOCK,
    ],
)

build(
    "list_of_shareholders.docx",
    "LIST OF SHAREHOLDERS",
    [
        "{{ company_name }} (CIN: {{ cin }}) — as on {{ document_date }}",
        "",
        "{% for s in shareholders %}",
        "{{ loop.index }}. {{ s.name }} — Folio: {{ s.folio }} — Shares: {{ s.shares }} "
        "({{ s.percentage }}%) — {{ s.category }}",
        "{% endfor %}",
        "",
        "Total shares: {{ total_shares }}",
        *SIGNING_BLOCK,
    ],
)


build(
    "board_notice.docx",
    "NOTICE OF MEETING OF THE BOARD OF DIRECTORS",
    [
        "NOTICE is hereby given that a {{ meeting_label }} of the Board of "
        "Directors of {{ company_name }} (CIN: {{ cin }}) will be held on "
        "{{ meeting_date }} at {{ meeting_time }} at {{ venue }}.",
        "",
        "AGENDA:",
        "{% for item in agenda_items %}",
        "{{ loop.index }}. {{ item }}",
        "{% endfor %}",
        "",
        "You are requested to make it convenient to attend the meeting.",
        *SIGNING_BLOCK,
    ],
)


build(
    "shorter_notice_consent.docx",
    "CONSENT FOR SHORTER NOTICE",
    [
        "To,",
        "The Board of Directors,",
        "{{ company_name }} (CIN: {{ cin }})",
        "{{ registered_address }}",
        "",
        "Subject: Consent to convene the {{ meeting_label }} at shorter notice "
        "(proviso to Section 101(1) of the Companies Act, 2013).",
        "",
        "I/We, {{ member_name }} (Folio: {{ folio_no }}), being a member of the "
        "company holding {{ shares_held }} shares, hereby accord consent to the "
        "convening of the {{ meeting_label }} of the company on {{ meeting_date }} "
        "at {{ meeting_time }} at {{ venue }} at notice shorter than that "
        "prescribed under Section 101(1) of the Companies Act, 2013.",
        "",
        "_______________________",
        "{{ member_name }}",
        "Folio: {{ folio_no }}",
        "Date: {{ document_date }}",
        "Place: {{ place }}",
    ],
)

build(
    "auditor_appointment.docx",
    "INTIMATION OF APPOINTMENT AS STATUTORY AUDITOR",
    [
        "To,",
        "{{ auditor_name }} (FRN: {{ auditor_frn }})",
        "{{ auditor_address }}",
        "",
        "Subject: Appointment as Statutory Auditor under Section 139 of the "
        "Companies Act, 2013.",
        "",
        "We are pleased to inform you that at the {{ meeting_label }} of "
        "{{ company_name }} (CIN: {{ cin }}) held on {{ meeting_date }}, your "
        "firm was appointed as the Statutory Auditor of the company, to hold "
        "office from the financial year ending {{ from_fy_label }} on the terms "
        "approved thereat.",
        "",
        "You are requested to confirm your eligibility under Section 141 of the "
        "Companies Act, 2013. The company shall file the notice of appointment "
        "in Form ADT-1 within the prescribed period.",
        *SIGNING_BLOCK,
    ],
)

build(
    "mr3_secretarial_audit.docx",
    "FORM No. MR-3 — SECRETARIAL AUDIT REPORT (skeleton)",
    [
        "For the financial year ended {{ period_ended }}",
        "",
        "To,",
        "The Members,",
        "{{ company_name }} (CIN: {{ cin }})",
        "{{ registered_address }}",
        "",
        "We have conducted the secretarial audit of the compliance of applicable "
        "statutory provisions and the adherence to good corporate practices by "
        "{{ company_name }} (hereinafter called \"the Company\"). The audit was "
        "conducted in a manner that provided us a reasonable basis for evaluating "
        "the corporate conducts / statutory compliances and expressing our "
        "opinion thereon.",
        "",
        "Based on our verification of the Company's books, papers, minute books, "
        "forms and returns filed and other records maintained by the Company, we "
        "hereby report that in our opinion, during the audit period covering the "
        "financial year ended {{ period_ended }}, the Company has complied with "
        "the statutory provisions listed hereunder:",
        "",
        "(i) The Companies Act, 2013 and the rules made thereunder;",
        "(ii) The Securities Contracts (Regulation) Act, 1956 and rules;",
        "(iii) The Depositories Act, 1996 and regulations;",
        "(iv) The Foreign Exchange Management Act, 1999 to the extent applicable;",
        "(v) Other laws specifically applicable to the Company as identified by "
        "the management: {{ other_applicable_laws }}",
        "",
        "OBSERVATIONS / QUALIFICATIONS:",
        "{{ observations }}",
        "",
        "_______________________",
        "{{ pcs_name }}",
        "Practising Company Secretary",
        "Membership No.: {{ pcs_membership_no }} · COP No.: {{ pcs_cop_no }}",
        "Date: {{ document_date }}",
        "Place: {{ place }}",
    ],
)

manifest = [
    {"code": "BOARD-NOTICE", "name": "Board Meeting Notice",
     "governing_reference": "S.173 / SS-1", "file": "board_notice.docx", "version": 1},
    {"code": "AGM-NOTICE", "name": "AGM Notice", "governing_reference": "S.101 / SS-2",
     "file": "agm_notice.docx", "version": 1},
    {"code": "DIRECTORS-REPORT", "name": "Directors' Report (skeleton)",
     "governing_reference": "S.134", "file": "directors_report.docx", "version": 1},
    {"code": "MEETING-MINUTES", "name": "Board/AGM Meeting Minutes",
     "governing_reference": "S.118 / SS-1, SS-2", "file": "meeting_minutes.docx", "version": 1},
    {"code": "ATTENDANCE-SHEET", "name": "Attendance Sheet",
     "governing_reference": "SS-1 / SS-2", "file": "attendance_sheet.docx", "version": 1},
    {"code": "LIST-DIRECTORS", "name": "List of Directors",
     "governing_reference": "S.170", "file": "list_of_directors.docx", "version": 1},
    {"code": "SHORTER-NOTICE", "name": "Consent for Shorter Notice",
     "governing_reference": "S.101(1) proviso / SS-2", "file": "shorter_notice_consent.docx", "version": 1},
    {"code": "AUDITOR-APPOINTMENT", "name": "Auditor Appointment Intimation",
     "governing_reference": "S.139 / S.141", "file": "auditor_appointment.docx", "version": 1},
    {"code": "MR-3", "name": "Secretarial Audit Report (skeleton)",
     "governing_reference": "S.204 / R.9 / Form MR-3", "file": "mr3_secretarial_audit.docx", "version": 1},
    {"code": "LIST-SHAREHOLDERS", "name": "List of Shareholders",
     "governing_reference": "S.88", "file": "list_of_shareholders.docx", "version": 1},
]
# validation stamps deliberately absent: templates are unusable until a named
# professional stamps them via PUT /api/v1/templates/{code}/validate
(OUT / "manifest.json").write_text(json.dumps(manifest, indent=2) + "\n")
print("built manifest.json (unstamped)")
